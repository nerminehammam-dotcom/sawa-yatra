import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


def paragraph_record(index, paragraph):
    numbering = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
    return {
        "index": index,
        "style": paragraph.style.name if paragraph.style else None,
        "text": paragraph.text,
        "numbered": numbering is not None,
        "page_break_before": bool(paragraph.paragraph_format.page_break_before),
        "keep_with_next": bool(paragraph.paragraph_format.keep_with_next),
        "run_count": len(paragraph.runs),
    }


def table_record(index, table):
    rows = []
    for row in table.rows:
        rows.append([cell.text for cell in row.cells])
    return {
        "index": index,
        "style": table.style.name if table.style else None,
        "rows": len(table.rows),
        "columns": len(table.columns),
        "data": rows,
    }


def main(path):
    source = Path(path)
    document = Document(source)

    sections = []
    for index, section in enumerate(document.sections):
        sections.append(
            {
                "index": index,
                "width_inches": section.page_width.inches,
                "height_inches": section.page_height.inches,
                "top_margin_inches": section.top_margin.inches,
                "right_margin_inches": section.right_margin.inches,
                "bottom_margin_inches": section.bottom_margin.inches,
                "left_margin_inches": section.left_margin.inches,
                "header_distance_inches": section.header_distance.inches,
                "footer_distance_inches": section.footer_distance.inches,
            }
        )

    paragraphs = [
        paragraph_record(index, paragraph)
        for index, paragraph in enumerate(document.paragraphs)
    ]
    tables = [table_record(index, table) for index, table in enumerate(document.tables)]
    blocks = []
    paragraph_index = 0
    table_index = 0
    for block_index, block in enumerate(document.iter_inner_content()):
        if isinstance(block, Paragraph):
            blocks.append(
                {
                    "block_index": block_index,
                    "type": "paragraph",
                    "paragraph_index": paragraph_index,
                    "record": paragraph_record(paragraph_index, block),
                }
            )
            paragraph_index += 1
        elif isinstance(block, Table):
            blocks.append(
                {
                    "block_index": block_index,
                    "type": "table",
                    "table_index": table_index,
                    "record": table_record(table_index, block),
                }
            )
            table_index += 1

    with zipfile.ZipFile(source) as archive:
        media = []
        for name in sorted(archive.namelist()):
            if name.startswith("word/media/"):
                info = archive.getinfo(name)
                media.append({"name": name, "bytes": info.file_size})
        has_comments = "word/comments.xml" in archive.namelist()
        has_numbering = "word/numbering.xml" in archive.namelist()

    result = {
        "source": str(source),
        "core_properties": {
            "title": document.core_properties.title,
            "subject": document.core_properties.subject,
            "author": document.core_properties.author,
            "keywords": document.core_properties.keywords,
        },
        "sections": sections,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "inline_shape_count": len(document.inline_shapes),
        "paragraphs": paragraphs,
        "tables": tables,
        "blocks": blocks,
        "media": media,
        "has_comments": has_comments,
        "has_numbering": has_numbering,
    }
    rendered = json.dumps(result, ensure_ascii=False, indent=2)
    if len(sys.argv) > 2:
        Path(sys.argv[2]).write_text(rendered, encoding="utf-8")
    else:
        print(rendered)


if __name__ == "__main__":
    main(sys.argv[1])
